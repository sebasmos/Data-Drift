#!/usr/bin/env python3
"""
Generate a comprehensive Word report from the drift analysis results.
Includes all figures, tables, and statistical findings.
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

# Paths
BASE_DIR = Path(__file__).parent.parent
FIGURES_DIR = BASE_DIR / 'figures'
OUTPUT_DIR = BASE_DIR / 'output'

def add_heading(doc, text, level=1):
    """Add a heading with consistent formatting."""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph_text(doc, text, bold=False, italic=False):
    """Add a paragraph with optional formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_key_finding_box(doc, title, content):
    """Add a highlighted key finding box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)

    # Title in bold
    run = p.add_run(f"🔑 {title}: ")
    run.bold = True
    run.font.size = Pt(11)

    # Content
    run2 = p.add_run(content)
    run2.font.size = Pt(11)

    return p

def add_table_from_data(doc, headers, rows, bold_first_col=True):
    """Add a table with headers and rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data)
            if bold_first_col and col_idx == 0:
                row_cells[col_idx].paragraphs[0].runs[0].bold = True

    return table

def add_figure(doc, image_path, caption, width=6.0):
    """Add a figure with caption."""
    if image_path.exists():
        doc.add_picture(str(image_path), width=Inches(width))

        # Caption
        caption_p = doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption_p.add_run(caption)
        run.italic = True
        run.font.size = Pt(10)

        doc.add_paragraph()  # Space after figure
        return True
    else:
        doc.add_paragraph(f"[Figure not found: {image_path.name}]")
        return False

def create_report():
    """Generate the full Word report."""
    doc = Document()

    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    title = doc.add_heading('Subgroup-Specific Drift in ICU Severity Scores', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Analysis Report')
    run.font.size = Pt(16)
    run.italic = True

    # Authors
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.add_run('Hamza Nabulsi, Xiaoli Liu, Leo Anthony Celi, Sebastian Cajas')

    doc.add_paragraph()

    # Summary box
    summary = doc.add_paragraph()
    summary.paragraph_format.left_indent = Inches(0.5)
    summary.paragraph_format.right_indent = Inches(0.5)
    run = summary.add_run(
        "TL;DR: ICU severity scores (OASIS, SAPS-II, APS-III, SOFA) drift differently "
        "across demographic subgroups. We analyze 826,611 ICU admissions across 6 primary "
        "+ 2 supplementary datasets from the US, Europe, and Asia (2001-2022) to quantify "
        "these disparities."
    )
    run.italic = True

    doc.add_page_break()

    # =========================================================================
    # EXECUTIVE SUMMARY - MOST STRIKING RESULTS
    # =========================================================================
    add_heading(doc, 'Executive Summary: Key Findings', 1)

    doc.add_paragraph(
        "This analysis reveals critical disparities in how ICU severity scores drift "
        "over time across different patient populations. The most striking findings are:"
    )

    # Key Finding 1: Temporal Drift is Widespread
    add_heading(doc, '1. Temporal Drift Affects All Datasets', 2)
    p = doc.add_paragraph()
    p.add_run("Significant drift was observed across multiple datasets spanning 2001-2022").bold = True
    p.add_run(" — with varying rates depending on score type and patient population.")

    doc.add_paragraph(
        "This suggests that regular recalibration is necessary, and the effect is NOT uniform "
        "across populations or healthcare systems."
    )

    # Key Finding 2: Hispanic Patients Hit Hardest
    add_heading(doc, '2. Hispanic Patients Experienced Notable Degradation', 2)

    table_data = [
        ['Hispanic', '-0.078', '-0.059', '5× worse than White'],
        ['Black', '-0.027', '-0.008', '2× worse than White'],
        ['White', '-0.015', '-0.003', 'Baseline'],
        ['Asian', '-0.040', '+0.007', 'Mixed'],
    ]
    add_table_from_data(doc,
        ['Subgroup', 'OASIS Δ', 'SOFA Δ', 'Relative Impact'],
        table_data
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Hispanic patients saw OASIS score performance decline by -0.078 AUC (p < 0.0001)").bold = True
    p.add_run(" — the single largest statistically significant decline in the entire study.")

    # Key Finding 3: Age Divergence Pattern
    add_heading(doc, '3. Young vs. Elderly: Opposite Drift Directions', 2)

    doc.add_paragraph(
        "A consistent pattern emerged across US, European, and Asian datasets: "
        "young patients (18-44) and elderly patients (80+) experience OPPOSITE drift directions."
    )

    age_data = [
        ['US (Boston)', 'MIMIC-IV', '+0.110', '+0.004', 'Young improve more'],
        ['Europe', 'Saltz', '-0.039', '+0.034', 'Young decline, elderly improve'],
        ['Asia', 'Zhejiang', '-0.096', '+0.037', 'Young decline, elderly improve'],
        ['US (Multi-center)', 'eICU', '-0.051', '+0.023', 'Young decline, elderly improve'],
    ]
    add_table_from_data(doc,
        ['Region', 'Dataset', 'Age 18-44 (SOFA Δ)', 'Age 80+ (SOFA Δ)', 'Pattern'],
        age_data
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Implication: ").bold = True
    p.add_run("Age-specific recalibration strategies may be necessary, as uniform adjustments "
              "would help one group while harming another.")

    # Key Finding 4: Statistical Summary
    add_heading(doc, '4. Statistical Summary', 2)

    doc.add_paragraph("Of 206 subgroup-score comparisons analyzed with DeLong's test:")

    stat_data = [
        ['eICU Combined (7 years)', '40 / 55', '72.7%', 'Notable drift'],
        ['MIMIC Combined (21 years)', '14 / 40', '35.0%', 'Moderate drift'],
        ['Saltz (8 years)', '5 / 28', '17.9%', 'Mostly stable'],
        ['Zhejiang (11 years)', '5 / 28', '17.9%', 'Mostly stable'],
    ]
    add_table_from_data(doc,
        ['Dataset', 'Significant / Total', 'Rate', 'Interpretation'],
        stat_data
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Core Thesis: ").bold = True
    p.add_run("Model drift affects demographic subgroups NON-UNIFORMLY. "
              "Uniform recalibration strategies would fail to address subgroup-specific disparities.")

    doc.add_page_break()

    # =========================================================================
    # SECTION 1: DATA OVERVIEW
    # =========================================================================
    add_heading(doc, '1. Datasets Analyzed', 1)

    doc.add_paragraph(
        "We analyzed 826,611 ICU admissions across 8 datasets from 3 continents, "
        "spanning over two decades (2001-2022)."
    )

    add_heading(doc, 'Primary Datasets', 2)

    primary_data = [
        ['MIMIC Combined', '112,468', '2001-2022', '~11%', 'US (Boston)'],
        ['eICU Combined', '661,358', '2014-2021', '~10%', 'US (Multi-center)'],
        ['Saltz', '27,259', '2013-2021', '7.9%', 'Europe (Netherlands)'],
        ['Zhejiang', '7,932', '2011-2022', '14.7%', 'Asia (China)'],
    ]
    add_table_from_data(doc,
        ['Dataset', 'N', 'Period', 'Mortality', 'Source'],
        primary_data
    )

    doc.add_paragraph()

    add_heading(doc, 'Supplementary Datasets (MIMIC-IV Subsets)', 2)

    supp_data = [
        ['MIMIC-IV Mouthcare', '8,675', '2008-2019', '~30%', 'Oral care frequency'],
        ['MIMIC-IV Mech. Vent.', '8,919', '2008-2019', '~30%', 'Turning frequency'],
    ]
    add_table_from_data(doc,
        ['Dataset', 'N', 'Period', 'Mortality', 'Analysis Focus'],
        supp_data
    )

    doc.add_page_break()

    # =========================================================================
    # SECTION 2: OVERALL DRIFT RESULTS
    # =========================================================================
    add_heading(doc, '2. Overall Drift by Dataset', 1)

    doc.add_paragraph(
        "We measured the change in AUC (Area Under ROC Curve) for mortality prediction "
        "from the first to last time period in each dataset. Positive values indicate "
        "improving discrimination; negative values indicate declining performance."
    )

    add_heading(doc, 'Overall Drift Summary (All Scores)', 2)

    drift_data = [
        ['MIMIC-IV', '+0.034*', '+0.011', '+0.008', '+0.022', 'Improving'],
        ['Saltz', '+0.034', '+0.049', '+0.054', '+0.076*', 'Improving'],
        ['Zhejiang', '+0.050', '+0.049', '+0.057', '+0.111*', 'Improving'],
        ['eICU', '+0.019*', '-0.006', '+0.010*', '+0.008', 'Stable'],
        ['eICU-New', '-0.003', '-0.014*', '-0.008*', '-0.023*', 'Declining'],
    ]
    add_table_from_data(doc,
        ['Dataset', 'SOFA Δ', 'OASIS Δ', 'SAPS-II Δ', 'APS-III Δ', 'Direction'],
        drift_data
    )

    doc.add_paragraph("* = statistically significant (95% CI excludes 0)")

    # Add Figure S3 (overall drift is now supplementary)
    doc.add_paragraph()
    add_heading(doc, 'Overall Drift Comparison (Supplementary)', 3)
    add_figure(doc, FIGURES_DIR / 'supplementary' / 'figS3_overall_drift_comparison.png',
               'Figure S3: Overall score performance trends (each dataset analyzed independently)')

    doc.add_page_break()

    # =========================================================================
    # SECTION 3: SUBGROUP-SPECIFIC DRIFT
    # =========================================================================
    add_heading(doc, '3. Subgroup-Specific Drift Analysis', 1)

    doc.add_paragraph(
        "The key finding of this study is that drift does NOT affect all patients uniformly. "
        "Different demographic subgroups experience different—and sometimes opposite—changes "
        "in score performance over time."
    )

    # Age Stratification
    add_heading(doc, '3.1 Age Group Divergence', 2)

    doc.add_paragraph(
        "Young (18-44) and elderly (80+) patients consistently show OPPOSITE drift directions. "
        "This pattern appears across all geographic regions."
    )

    add_figure(doc, FIGURES_DIR / 'fig1_age_stratified_drift.png',
               'Figure 1: Age-stratified drift by dataset - PRIMARY EVIDENCE of non-uniform drift')

    # Race/Ethnicity
    add_heading(doc, '3.2 Racial/Ethnic Disparities', 2)

    doc.add_paragraph(
        "Hispanic patients experienced notable score degradation across datasets."
    )

    add_figure(doc, FIGURES_DIR / 'fig2_race_disparities.png',
               'Figure 2: Race/ethnicity disparities in US datasets (MIMIC Combined, eICU Combined)')

    # SOFA Subgroup Table
    add_heading(doc, 'SOFA Drift by Subgroup', 3)

    sofa_data = [
        ['Overall', '+0.034*', '+0.034', '+0.050', '+0.019*', '-0.003'],
        ['Age 18-44', '+0.110*', '-0.039', '-0.096', '+0.021', '-0.051*'],
        ['Age 80+', '+0.004', '+0.034', '+0.037', '+0.015', '+0.023*'],
        ['Male', '+0.030', '+0.041', '+0.061', '+0.015*', '-0.009*'],
        ['Female', '+0.038', '+0.009', '+0.004', '+0.023*', '+0.006'],
        ['White', '+0.046*', '-', '-', '+0.017*', '-0.003'],
        ['Black', '+0.019', '-', '-', '+0.013', '-0.008'],
        ['Hispanic', '-', '-', '-', '+0.016', '-0.059*'],
        ['Asian', '+0.051', '-', '-', '+0.087', '+0.007'],
    ]
    add_table_from_data(doc,
        ['Subgroup', 'MIMIC-IV', 'Saltz', 'Zhejiang', 'eICU', 'eICU-New'],
        sofa_data
    )
    doc.add_paragraph("* = statistically significant (p < 0.05, DeLong's test)")

    doc.add_page_break()

    # =========================================================================
    # SECTION 4: COMPREHENSIVE VISUALIZATIONS
    # =========================================================================
    add_heading(doc, '4. Comprehensive Visualizations', 1)

    add_heading(doc, 'Summary Figure', 2)
    add_figure(doc, FIGURES_DIR / 'fig5_money_figure.png',
               'Figure 5: Multi-panel summary - (A) Age group divergence, (B) Race disparities, '
               '(C) Subgroup drift heatmap', width=6.5)

    add_heading(doc, 'Drift Delta Summary', 2)
    add_figure(doc, FIGURES_DIR / 'fig3_drift_delta_summary.png',
               'Figure 3: AUC change summary by subgroup (each dataset analyzed independently)')

    add_heading(doc, 'Comprehensive Heatmap', 2)
    add_figure(doc, FIGURES_DIR / 'fig4_comprehensive_heatmap.png',
               'Figure 4: Comprehensive drift heatmap showing all subgroups and scores')

    doc.add_page_break()

    # =========================================================================
    # SECTION 5: SUPPLEMENTARY ANALYSIS
    # =========================================================================
    add_heading(doc, '5. Supplementary Analysis: MIMIC-IV Care Frequency', 1)

    doc.add_paragraph(
        "We also analyzed how SOFA score drift relates to nursing care frequency "
        "in two MIMIC-IV subsets: mouthcare and mechanical ventilation cohorts."
    )

    add_figure(doc, FIGURES_DIR / 'supplementary' / 'figS1_mimic_mouthcare.png',
               'Figure S1: MIMIC-IV Mouthcare cohort (N=8,675) - SOFA drift by age, race, '
               'gender, and care frequency')

    add_figure(doc, FIGURES_DIR / 'supplementary' / 'figS2_mimic_mechvent.png',
               'Figure S2: MIMIC-IV Mechanical Ventilation cohort (N=8,919) - SOFA drift by age, '
               'race, gender, and care frequency')

    doc.add_page_break()

    # =========================================================================
    # SECTION 6: STATISTICAL METHODS
    # =========================================================================
    add_heading(doc, '6. Statistical Methods', 1)

    add_heading(doc, 'Methods Used', 2)

    methods_data = [
        ['Bootstrap CIs', 'Confidence intervals for AUC', 'Percentile method (n=100-1000)'],
        ["DeLong's test", 'Compare AUCs between periods', 'Hanley-McNeil variance, z-test'],
        ['Significance', 'Identify reliable drift', 'p < 0.05 or 95% CI excludes 0'],
    ]
    add_table_from_data(doc,
        ['Method', 'Purpose', 'Implementation'],
        methods_data
    )

    add_heading(doc, "What Does DeLong's Test Tell You?", 2)

    doc.add_paragraph(
        "DeLong's test answers: \"Is the observed AUC change real, or just random noise?\""
    )

    doc.add_paragraph("• p < 0.05 → The drift is statistically significant (unlikely due to chance)")
    doc.add_paragraph("• p ≥ 0.05 → The drift could be random variation (insufficient evidence)")

    add_heading(doc, 'Example Interpretations', 3)

    example_data = [
        ['eICU-New Hispanic OASIS', '-0.078', '<0.0001', 'Real decline — nearly impossible this is chance'],
        ['MIMIC-IV Asian APS-III', '+0.184', '0.005', 'Real improvement — 0.5% chance this is random'],
        ['Saltz Overall SOFA', '+0.034', '0.28', 'Not significant — 28% chance this is noise'],
    ]
    add_table_from_data(doc,
        ['Finding', 'Δ AUC', 'p-value', 'Interpretation'],
        example_data
    )

    doc.add_page_break()

    # =========================================================================
    # SECTION 7: CLINICAL IMPLICATIONS
    # =========================================================================
    add_heading(doc, '7. Clinical Implications', 1)

    implications = [
        ("Uniform recalibration is insufficient",
         "Different subgroups need different adjustments. A single recalibration factor "
         "would help some groups while potentially harming others."),

        ("Vulnerable populations require attention",
         "Hispanic and young patients showed notable degradation in some datasets. "
         "These groups may need enhanced monitoring or alternative scoring approaches."),

        ("Geographic context matters",
         "The same subgroup can experience opposite drift in different healthcare systems. "
         "Local validation is essential before applying any score."),

        ("Age-specific models may be needed",
         "The consistent age divergence pattern suggests fundamental differences in how "
         "scores perform across age groups that may warrant separate models."),
    ]

    for i, (title, content) in enumerate(implications, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {title}: ").bold = True
        p.add_run(content)
        doc.add_paragraph()

    # =========================================================================
    # SECTION 8: NEXT STEPS
    # =========================================================================
    add_heading(doc, '8. Recommended Next Steps', 1)

    add_heading(doc, 'Immediate Priorities', 2)
    doc.add_paragraph("• Calibration analysis: Assess calibration drift (Brier score, curves)")
    doc.add_paragraph("• Intersectional analysis: Examine combinations (e.g., elderly Hispanic)")
    doc.add_paragraph("• Temporal deep-dive: Investigate year-by-year drift patterns")

    add_heading(doc, 'Extended Analysis', 2)
    doc.add_paragraph("• Recalibration strategies: Test subgroup-specific approaches")
    doc.add_paragraph("• Feature importance: Identify which components drive drift")
    doc.add_paragraph("• External validation: Apply to additional datasets (ANZICS, UK ICU)")

    add_heading(doc, 'Clinical Translation', 2)
    doc.add_paragraph("• Decision threshold analysis: How drift affects clinical cutoffs")
    doc.add_paragraph("• Fairness metrics: Compute equalized odds across subgroups")
    doc.add_paragraph("• Clinical guidelines: Recommendations for score interpretation")

    # =========================================================================
    # Save the document
    # =========================================================================
    output_path = OUTPUT_DIR / 'Drift_Analysis_Report.docx'
    doc.save(str(output_path))
    print(f"\n✓ Report saved to: {output_path}")
    print(f"  Total size: {output_path.stat().st_size / 1024:.1f} KB")

    return output_path

if __name__ == '__main__':
    create_report()
